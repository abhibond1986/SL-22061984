"""
test_service.py — self-checks for the Safety Lens OCR service.

Split into two tiers so it is useful even before the heavy install:

    Tier 1 (no PaddleOCR needed): text cleaning, clause chunking, DOCX
           reading, PDF text-layer reading.
    Tier 2 (needs paddleocr):     real OCR of a generated image.

Run everything:      python test_service.py
Skip OCR:            python test_service.py --no-ocr

Exit code is non-zero if any check fails, so this can gate a deploy.
"""

from __future__ import annotations

import io
import sys
import traceback

PASS, FAIL = "PASS", "FAIL"
results: list[tuple[str, str, str]] = []


def check(name: str, fn) -> None:
    try:
        detail = fn() or ""
        results.append((PASS, name, str(detail)))
        print(f"  [{PASS}] {name} {detail}")
    except Exception as exc:
        results.append((FAIL, name, f"{type(exc).__name__}: {exc}"))
        print(f"  [{FAIL}] {name} -> {type(exc).__name__}: {exc}")
        traceback.print_exc(limit=2)


# --------------------------------------------------------------------------
# Tier 1: cleaning
# --------------------------------------------------------------------------

def test_clean_dehyphenates():
    from extractors import clean_text
    out = clean_text("The equip-\nment must be isolated.")
    assert "equipment" in out, out
    return "-> 'equipment'"


def test_clean_strips_page_numbers():
    from extractors import clean_text
    out = clean_text("Wear PPE.\nPage 4 of 12\nIsolate the supply.")
    assert "Page 4 of 12" not in out, out
    assert "Isolate the supply." in out
    return "bare page number removed"


def test_clean_collapses_blank_runs():
    from extractors import clean_text
    out = clean_text("A\n\n\n\n\nB")
    assert out == "A\n\nB", repr(out)
    return "5 newlines -> 2"


def test_clean_removes_control_chars():
    from extractors import clean_text
    out = clean_text("Hot\x00 metal\x07 zone")
    assert "\x00" not in out and "\x07" not in out
    return "NUL/BEL stripped"


def test_clean_handles_empty():
    from extractors import clean_text
    assert clean_text("") == ""
    assert clean_text(None) == ""
    return "empty + None safe"


# --------------------------------------------------------------------------
# Tier 1: chunking
# --------------------------------------------------------------------------

SOP_TEXT = """## SOP-4412 LADLE HANDLING

1. PURPOSE
This procedure defines the controls for handling hot metal ladles in the steel
melting shop, covering transfer, teeming and parking operations.

2. SCOPE
Applies to all SMS-1 and SMS-2 crane and ladle operations, including
contractor personnel working under a valid work permit.

4.2 Personal Protective Equipment
Operators shall wear aluminised proximity suits, a face shield rated to the
radiant heat load, and heat-resistant gloves before approaching the ladle.
Cotton undergarments are mandatory; synthetic fabric is prohibited because it
melts onto skin under radiant heat.

4.2.1 Lockout Tagout
Before any maintenance on the ladle turret, isolate the drive at the local
isolator, apply a personal lock, and verify zero energy by attempting a start.
Only the person who applied the lock may remove it.

4.3 Emergency Response
In the event of a hot metal spill, sound the evacuation alarm, withdraw upwind
along the marked route, and do not attempt to apply water to spilled metal.
"""


def _chunks():
    from chunker import chunk_pages
    return chunk_pages([{"page": 1, "text": SOP_TEXT}])


def test_chunker_produces_chunks():
    cs = _chunks()
    assert len(cs) >= 3, f"only {len(cs)} chunks"
    return f"{len(cs)} chunks"


def test_chunker_captures_clause_numbers():
    cs = _chunks()
    found = {c.clause_no for c in cs if c.clause_no}
    assert "4.2" in found, f"clause 4.2 missing from {sorted(found)}"
    assert "4.2.1" in found, f"clause 4.2.1 missing from {sorted(found)}"
    return f"clauses {sorted(found)}"


def test_chunker_does_not_split_a_clause_across_chunks():
    """The LOTO rule must live in exactly one chunk, or citations mislead."""
    cs = _chunks()
    hits = [c for c in cs if "may remove it" in c.text]
    assert len(hits) == 1, f"{len(hits)} chunks contain the LOTO sentence"
    assert hits[0].clause_no == "4.2.1", f"attributed to {hits[0].clause_no}"
    return "LOTO text -> clause 4.2.1 only"


def test_chunker_tracks_pages():
    from chunker import chunk_pages
    cs = chunk_pages([
        {"page": 1, "text": "1. PURPOSE\n" + "Isolate the drive. " * 40},
        {"page": 2, "text": "2. SCOPE\n" + "Applies to SMS-1. " * 40},
    ])
    pages = {c.page_from for c in cs}
    assert pages == {1, 2}, f"pages seen: {pages}"
    return "page_from tracked for both pages"


def test_chunker_respects_target_size():
    from chunker import chunk_pages, TARGET_CHARS
    long_clause = "5.1 Controls\n" + ("The permit must be countersigned. " * 300)
    cs = chunk_pages([{"page": 1, "text": long_clause}])
    oversized = [c for c in cs if len(c.text) > TARGET_CHARS * 1.6]
    assert not oversized, f"{len(oversized)} chunks exceed the size budget"
    assert len(cs) > 1, "a 10k-char clause should have been split"
    return f"{len(cs)} pieces, max {max(len(c.text) for c in cs)} chars"


def test_chunker_ignores_decimals_in_prose():
    """'a 2.5 m platform' must not be mistaken for clause 2.5."""
    from chunker import chunk_pages
    cs = chunk_pages([{"page": 1,
                       "text": "Guard rails are required on any platform "
                               "above 2.5 m in height as per the standard."}])
    assert all(c.clause_no is None for c in cs), \
        f"false clause detected: {[c.clause_no for c in cs]}"
    return "mid-sentence decimal not treated as a clause"


def test_chunker_empty_input():
    from chunker import chunk_pages
    assert chunk_pages([]) == []
    assert chunk_pages([{"page": 1, "text": ""}]) == []
    return "empty input safe"


def test_chunker_no_text_is_lost():
    """Every alphanumeric word in the source must survive into some chunk."""
    import re
    cs = _chunks()
    joined = " ".join(c.text for c in cs)
    src_words = set(re.findall(r"[A-Za-z]{4,}", SOP_TEXT))
    got_words = set(re.findall(r"[A-Za-z]{4,}", joined))
    missing = src_words - got_words
    assert not missing, f"lost words: {sorted(missing)[:10]}"
    return f"all {len(src_words)} distinct words retained"


# --------------------------------------------------------------------------
# Tier 1: DOCX
# --------------------------------------------------------------------------

def test_docx_reads_paragraphs_and_tables():
    import docx
    from extractors import extract_docx

    d = docx.Document()
    d.add_heading("SOP-9001 Crane Operation", level=1)
    d.add_paragraph("Only licensed operators may operate the EOT crane.")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Hazard"
    table.cell(0, 1).text = "Control"
    table.cell(1, 0).text = "Load swing"
    table.cell(1, 1).text = "Use tag line"
    buf = io.BytesIO()
    d.save(buf)

    res = extract_docx(buf.getvalue())
    text = res.text
    assert "licensed operators" in text, "paragraph lost"
    # Tables carry the hazard/control matrix — losing them is the big risk.
    assert "Load swing" in text, "table cell lost"
    assert "Use tag line" in text, "table cell lost"
    assert res.pages[0].method == "docx"
    return f"{len(text)} chars, table preserved"


def test_docx_rejects_garbage():
    from extractors import extract_docx, UnsupportedDocument
    try:
        extract_docx(b"this is definitely not a docx")
    except UnsupportedDocument as exc:
        assert "Save As" in str(exc) or "Word" in str(exc)
        return "actionable error returned"
    raise AssertionError("garbage input should have raised UnsupportedDocument")


# --------------------------------------------------------------------------
# Tier 1: PDF text layer
# --------------------------------------------------------------------------

def test_pdf_text_layer_path():
    """A generated text PDF must be read WITHOUT invoking OCR."""
    try:
        from reportlab.pdfgen import canvas
    except ImportError:
        return "SKIPPED (reportlab not installed)"

    from extractors import extract_pdf

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    y = 800
    for line in [
        "SOP-4412 LADLE HANDLING",
        "4.2 Personal Protective Equipment",
        "Operators shall wear aluminised proximity suits and a face shield",
        "rated to the radiant heat load before approaching the ladle.",
        "Cotton undergarments are mandatory; synthetic fabric is prohibited.",
    ]:
        c.drawString(60, y, line)
        y -= 20
    c.showPage()
    c.save()

    res = extract_pdf(buf.getvalue())
    assert res.pages, "no pages returned"
    assert res.pages[0].method == "text-layer", \
        f"expected text-layer, got {res.pages[0].method} (OCR ran needlessly)"
    assert "proximity suits" in res.text, res.text[:200]
    return "read via text layer, OCR skipped"


def test_pdf_corrupt_does_not_crash():
    from extractors import extract_pdf
    res = extract_pdf(b"%PDF-1.4 truncated garbage")
    assert isinstance(res.pages, list)
    return "corrupt PDF handled gracefully"


# --------------------------------------------------------------------------
# Tier 2: real OCR
# --------------------------------------------------------------------------

def test_ocr_reads_generated_image():
    from PIL import Image, ImageDraw
    from extractors import ocr_image_bytes

    img = Image.new("RGB", (1000, 260), "white")
    draw = ImageDraw.Draw(img)
    # Default bitmap font is small; scale up so PaddleOCR has real glyphs.
    draw.text((30, 40), "SAFETY FIRST", fill="black")
    draw.text((30, 90), "WEAR YOUR HELMET", fill="black")
    big = img.resize((3000, 780), Image.LANCZOS)

    buf = io.BytesIO()
    big.save(buf, format="PNG")

    text, conf = ocr_image_bytes(buf.getvalue())
    upper = (text or "").upper()
    assert "SAFETY" in upper or "HELMET" in upper, f"OCR read: {text!r}"
    return f"read {text!r} (conf={conf})"


# --------------------------------------------------------------------------
# Runner
# --------------------------------------------------------------------------

def main() -> int:
    run_ocr = "--no-ocr" not in sys.argv

    print("\nTier 1 — text cleaning")
    check("clean: de-hyphenates line breaks", test_clean_dehyphenates)
    check("clean: strips bare page numbers", test_clean_strips_page_numbers)
    check("clean: collapses blank-line runs", test_clean_collapses_blank_runs)
    check("clean: removes control characters", test_clean_removes_control_chars)
    check("clean: handles empty input", test_clean_handles_empty)

    print("\nTier 1 — clause chunking")
    check("chunker: produces chunks", test_chunker_produces_chunks)
    check("chunker: captures clause numbers", test_chunker_captures_clause_numbers)
    check("chunker: keeps a clause intact", test_chunker_does_not_split_a_clause_across_chunks)
    check("chunker: tracks page numbers", test_chunker_tracks_pages)
    check("chunker: respects size budget", test_chunker_respects_target_size)
    check("chunker: ignores prose decimals", test_chunker_ignores_decimals_in_prose)
    check("chunker: empty input safe", test_chunker_empty_input)
    check("chunker: loses no text", test_chunker_no_text_is_lost)

    print("\nTier 1 — document readers")
    check("docx: paragraphs + tables", test_docx_reads_paragraphs_and_tables)
    check("docx: rejects garbage clearly", test_docx_rejects_garbage)
    check("pdf: uses text layer, not OCR", test_pdf_text_layer_path)
    check("pdf: corrupt file safe", test_pdf_corrupt_does_not_crash)

    if run_ocr:
        print("\nTier 2 — PaddleOCR (slow: loads the model)")
        check("ocr: reads a generated image", test_ocr_reads_generated_image)
    else:
        print("\nTier 2 — skipped (--no-ocr)")

    failed = [r for r in results if r[0] == FAIL]
    passed = len(results) - len(failed)
    print(f"\n{'='*60}\n{passed} passed, {len(failed)} failed\n{'='*60}")
    if failed:
        for _, name, detail in failed:
            print(f"  FAIL  {name}: {detail}")
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
