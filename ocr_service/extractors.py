"""
extractors.py — Safety Lens OCR service: document text extraction.

Handles the three input kinds from the Safety Lens ingest flow:

    PDF   -> embedded text layer if present, else rasterise + PaddleOCR
    DOCX  -> python-docx (paragraphs + tables + headers)
    IMAGE -> PaddleOCR directly

Design notes
------------
* PaddleOCR is loaded LAZILY and cached process-wide. Constructing it costs
  several seconds and ~400 MB RSS, so we never build it for a DOCX or a
  text-layer PDF that does not need it.
* Licence hygiene: we use pdfplumber (MIT) for text extraction and pypdfium2
  (Apache/BSD) for rasterising, deliberately NOT PyMuPDF, which is AGPL and
  would encumber the wider Safety Lens codebase.
* Every page result records WHICH method produced it, so the Flutter UI can
  tell the user "3 pages read directly, 2 pages OCR'd" and so low-confidence
  OCR pages can be flagged as unverified (mirroring the existing
  `sop_scan` / `sop_scan_raw` trusted-vs-unverified split in the KB).
"""

from __future__ import annotations

import io
import logging
import re
import threading
from dataclasses import dataclass, field
from typing import List, Optional, Tuple

log = logging.getLogger("safetylens.ocr")

# --------------------------------------------------------------------------
# Tunables
# --------------------------------------------------------------------------

# A PDF page whose embedded text layer yields fewer than this many characters
# is treated as a scan and sent to OCR. Steel-plant SOPs are dense; a genuine
# text page essentially always clears 100 chars. Pages that are mostly a
# diagram legitimately fall through to OCR, which is what we want.
MIN_TEXT_LAYER_CHARS = 100

# Rasterising DPI for OCR. 200 is the sweet spot for PaddleOCR on scanned
# A4 SOPs: 150 loses small clause numbers, 300 triples latency for no gain.
RASTER_DPI = 200

# Hard cap so one 400-page manual cannot pin the free-tier CPU for an hour.
MAX_OCR_PAGES = 60

# PaddleOCR drops boxes below this confidence. Steel-plant scans are often
# poor (faded photocopies, stamps), so this is deliberately permissive; the
# per-page mean confidence is returned so the client can flag weak pages.
MIN_BOX_CONFIDENCE = 0.50


# --------------------------------------------------------------------------
# Result types
# --------------------------------------------------------------------------

@dataclass
class PageResult:
    page: int
    text: str
    method: str                      # "text-layer" | "paddleocr" | "docx"
    confidence: Optional[float] = None   # mean OCR box confidence, None if N/A


@dataclass
class ExtractResult:
    kind: str                        # "pdf" | "docx" | "image" | "text"
    pages: List[PageResult] = field(default_factory=list)
    truncated: bool = False          # True if MAX_OCR_PAGES clipped the doc

    @property
    def text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text.strip())

    @property
    def ocr_page_count(self) -> int:
        return sum(1 for p in self.pages if p.method == "paddleocr")

    @property
    def mean_confidence(self) -> Optional[float]:
        vals = [p.confidence for p in self.pages if p.confidence is not None]
        return round(sum(vals) / len(vals), 4) if vals else None


class UnsupportedDocument(Exception):
    """Raised for a file kind we deliberately do not handle (e.g. legacy .doc)."""


# --------------------------------------------------------------------------
# Lazy PaddleOCR singleton
# --------------------------------------------------------------------------

_ocr_lock = threading.Lock()
_ocr_cache: dict = {}


def get_paddle(lang: str = "en"):
    """
    Return a cached PaddleOCR instance for `lang`.

    Thread-safe and lazy: the first call pays the model-load cost, later calls
    are free. Cached per-language because a bilingual plant may need both
    English and Devanagari models resident.
    """
    lang = _normalise_lang(lang)
    if lang in _ocr_cache:
        return _ocr_cache[lang]

    with _ocr_lock:
        if lang in _ocr_cache:          # re-check inside the lock
            return _ocr_cache[lang]

        from paddleocr import PaddleOCR   # imported here: heavy

        log.info("Loading PaddleOCR model for lang=%s (first use)", lang)
        engine = PaddleOCR(
            use_angle_cls=True,   # rotated/skewed scans are common
            lang=lang,
            show_log=False,
            use_gpu=False,        # free hosting tiers are CPU-only
        )
        _ocr_cache[lang] = engine
        log.info("PaddleOCR model for lang=%s ready", lang)
        return engine


def _normalise_lang(lang: str) -> str:
    """Map friendly language codes onto PaddleOCR's model names."""
    lang = (lang or "en").strip().lower()
    # PaddleOCR ships one shared Devanagari model; 'hi' is not a valid key.
    if lang in ("hi", "hin", "devanagari", "mr", "ne"):
        return "devanagari"
    if lang in ("en", "eng", "english"):
        return "en"
    return lang


# --------------------------------------------------------------------------
# OCR of a single raster image
# --------------------------------------------------------------------------

def ocr_image_bytes(data: bytes, lang: str = "en") -> Tuple[str, Optional[float]]:
    """
    OCR raw image bytes. Returns (text, mean_confidence).

    PaddleOCR's return shape differs between the 2.x and 3.x lines, so we
    normalise defensively rather than trusting one layout.
    """
    import numpy as np
    from PIL import Image, ImageOps

    try:
        img = Image.open(io.BytesIO(data))
    except Exception as exc:
        raise UnsupportedDocument(f"Not a readable image: {exc}") from exc

    # EXIF-rotate before OCR: phone photos of a noticeboard are usually the
    # worst-quality input this service sees, and a sideways image reads as
    # near-gibberish even with use_angle_cls.
    img = ImageOps.exif_transpose(img)
    img = img.convert("RGB")

    return _run_paddle(np.array(img), lang)


def _run_paddle(np_image, lang: str) -> Tuple[str, Optional[float]]:
    """Run PaddleOCR on an HxWx3 numpy array and flatten to reading order."""
    engine = get_paddle(lang)

    try:
        raw = engine.ocr(np_image, cls=True)
    except TypeError:
        # PaddleOCR 3.x dropped the `cls` kwarg.
        raw = engine.ocr(np_image)

    lines = _flatten_paddle_result(raw)
    if not lines:
        return "", None

    # Sort into human reading order. PaddleOCR is usually already top-to-bottom
    # but multi-column SOP layouts can interleave; bucketing y into ~12px bands
    # then sorting by x within a band recovers sane line order without a full
    # layout-analysis pass.
    lines.sort(key=lambda t: (round(t[0] / 12.0), t[1]))

    text = "\n".join(t[2] for t in lines)
    confs = [t[3] for t in lines if t[3] is not None]
    mean_conf = round(sum(confs) / len(confs), 4) if confs else None
    return text, mean_conf


def _flatten_paddle_result(raw) -> List[Tuple[float, float, str, Optional[float]]]:
    """
    Normalise PaddleOCR output into [(y, x, text, confidence), ...].

    Tolerates the three shapes seen in the wild:
      2.x:  [[ [box, (text, conf)], ... ]]        # outer list = images
      2.x:  [ [box, (text, conf)], ... ]          # single image, unwrapped
      3.x:  [{'rec_texts': [...], 'rec_scores': [...], 'rec_polys': [...]}]
    """
    out: List[Tuple[float, float, str, Optional[float]]] = []
    if not raw:
        return out

    # --- 3.x dict form ---
    first = raw[0] if isinstance(raw, (list, tuple)) and raw else None
    if isinstance(first, dict) and "rec_texts" in first:
        texts = first.get("rec_texts") or []
        scores = first.get("rec_scores") or []
        polys = first.get("rec_polys") or first.get("dt_polys") or []
        for i, txt in enumerate(texts):
            conf = float(scores[i]) if i < len(scores) else None
            if conf is not None and conf < MIN_BOX_CONFIDENCE:
                continue
            y = x = 0.0
            if i < len(polys) and polys[i] is not None:
                pts = polys[i]
                try:
                    y = float(min(p[1] for p in pts))
                    x = float(min(p[0] for p in pts))
                except Exception:
                    pass
            if str(txt).strip():
                out.append((y, x, str(txt).strip(), conf))
        return out

    # --- 2.x list forms ---
    # Unwrap the per-image nesting if present.
    blocks = raw
    if (
        isinstance(raw, (list, tuple))
        and len(raw) == 1
        and isinstance(raw[0], (list, tuple))
        and raw[0]
        and isinstance(raw[0][0], (list, tuple))
        and len(raw[0][0]) == 2
    ):
        blocks = raw[0]

    if blocks is None:
        return out

    for item in blocks:
        if not item or len(item) < 2:
            continue
        box, payload = item[0], item[1]
        if isinstance(payload, (list, tuple)) and len(payload) >= 2:
            txt, conf = str(payload[0]), float(payload[1])
        else:
            txt, conf = str(payload), None
        if conf is not None and conf < MIN_BOX_CONFIDENCE:
            continue
        if not txt.strip():
            continue
        y = x = 0.0
        try:
            y = float(min(p[1] for p in box))
            x = float(min(p[0] for p in box))
        except Exception:
            pass
        out.append((y, x, txt.strip(), conf))
    return out


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------

def extract_pdf(data: bytes, lang: str = "en", force_ocr: bool = False) -> ExtractResult:
    """
    Extract a PDF page by page, preferring the embedded text layer.

    Digital SOPs exported from Word have a perfect text layer — OCR'ing them
    would be slower AND less accurate. Scanned SOPs have no text layer at all.
    Mixed documents (a digital SOP with scanned annexures) are common in steel
    plants, so the decision is made PER PAGE, not per document.
    """
    result = ExtractResult(kind="pdf")

    text_by_page: List[str] = []
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                try:
                    text_by_page.append(page.extract_text() or "")
                except Exception as exc:
                    log.warning("text-layer read failed on a page: %s", exc)
                    text_by_page.append("")
    except Exception as exc:
        # A corrupt or encrypted PDF still might rasterise, so don't give up.
        log.warning("pdfplumber could not open the PDF (%s); OCR-only path", exc)
        text_by_page = []

    needs_ocr = [
        i for i, t in enumerate(text_by_page)
        if force_ocr or len(t.strip()) < MIN_TEXT_LAYER_CHARS
    ]

    # No text layer at all -> we don't yet know the page count; ask pypdfium2.
    rendered: dict = {}
    if not text_by_page or needs_ocr:
        rendered = _render_pdf_pages(
            data,
            page_indices=needs_ocr if text_by_page else None,
        )
        if not text_by_page:
            needs_ocr = sorted(rendered.keys())
            text_by_page = [""] * (max(needs_ocr) + 1 if needs_ocr else 0)

    if len(needs_ocr) > MAX_OCR_PAGES:
        log.warning(
            "Document needs OCR on %d pages; clipping to %d",
            len(needs_ocr), MAX_OCR_PAGES,
        )
        needs_ocr = needs_ocr[:MAX_OCR_PAGES]
        result.truncated = True

    ocr_set = set(needs_ocr)

    for i, layer_text in enumerate(text_by_page):
        if i in ocr_set and i in rendered:
            try:
                import numpy as np
                text, conf = _run_paddle(np.array(rendered[i]), lang)
                result.pages.append(
                    PageResult(page=i + 1, text=clean_text(text),
                               method="paddleocr", confidence=conf)
                )
                continue
            except Exception as exc:
                log.warning("OCR failed on page %d: %s", i + 1, exc)
                # Fall through and keep whatever thin text layer we had.

        result.pages.append(
            PageResult(page=i + 1, text=clean_text(layer_text), method="text-layer")
        )

    return result


def _render_pdf_pages(data: bytes, page_indices: Optional[List[int]] = None) -> dict:
    """Rasterise selected PDF pages to PIL images keyed by 0-based index."""
    out: dict = {}
    try:
        import pypdfium2 as pdfium
    except ImportError:
        log.error("pypdfium2 not installed — cannot rasterise scanned PDFs")
        return out

    scale = RASTER_DPI / 72.0
    doc = None
    try:
        doc = pdfium.PdfDocument(data)
        total = len(doc)
        targets = range(total) if page_indices is None else [
            i for i in page_indices if 0 <= i < total
        ]
        # Respect the OCR budget while rendering, not after — rendering 400
        # pages at 200 DPI is itself enough to OOM a 512 MB free-tier box.
        for i in list(targets)[:MAX_OCR_PAGES]:
            try:
                out[i] = doc[i].render(scale=scale).to_pil().convert("RGB")
            except Exception as exc:
                log.warning("Could not render page %d: %s", i + 1, exc)
    except Exception as exc:
        log.error("pypdfium2 failed to open the PDF: %s", exc)
    finally:
        if doc is not None:
            try:
                doc.close()
            except Exception:
                pass
    return out


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------

def extract_docx(data: bytes) -> ExtractResult:
    """
    Extract a .docx, including tables.

    Tables matter more than usual here: SOP hazard/control matrices and PPE
    requirement grids are almost always tables, and a paragraphs-only reader
    silently drops the single most safety-relevant part of the document.
    """
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise UnsupportedDocument("python-docx is not installed") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise UnsupportedDocument(
            "Could not read that Word file. If it is an older .doc, re-save "
            "it as .docx and try again."
        ) from exc

    parts: List[str] = []

    for para in document.paragraphs:
        txt = (para.text or "").strip()
        if not txt:
            continue
        style = (para.style.name or "") if para.style is not None else ""
        # Mark headings so the chunker can split on real section boundaries.
        if style.startswith("Heading") or style == "Title":
            parts.append(f"\n## {txt}")
        else:
            parts.append(txt)

    for t_index, table in enumerate(document.tables, start=1):
        rows: List[str] = []
        for row in table.rows:
            cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
            # Merged cells repeat their text across the span; collapse runs so
            # "PPE | PPE | PPE" doesn't pollute the retrieval text.
            collapsed: List[str] = []
            for cell in cells:
                if not collapsed or collapsed[-1] != cell:
                    collapsed.append(cell)
            if any(collapsed):
                rows.append(" | ".join(collapsed))
        if rows:
            parts.append(f"\n## Table {t_index}\n" + "\n".join(rows))

    result = ExtractResult(kind="docx")
    result.pages.append(
        PageResult(page=1, text=clean_text("\n".join(parts)), method="docx")
    )
    return result


# --------------------------------------------------------------------------
# Cleaning
# --------------------------------------------------------------------------

_RE_HYPHEN_BREAK = re.compile(r"(\w)-\n(\w)")
_RE_MANY_BLANKS = re.compile(r"\n{3,}")
_RE_TRAIL_SPACE = re.compile(r"[ \t]+$", re.MULTILINE)
_RE_MANY_SPACES = re.compile(r"[ \t]{2,}")
_RE_PAGE_NUM = re.compile(r"^\s*(?:page\s*)?\d+\s*(?:of\s*\d+)?\s*$", re.IGNORECASE)
_RE_BAD_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def clean_text(text: str) -> str:
    """
    Normalise extracted text for retrieval.

    Kept deliberately conservative — this text is quoted back to a safety
    officer as evidence, so we fix mechanical artefacts (control characters,
    hyphenated line breaks, bare page numbers) and nothing more. We do NOT
    fix spelling or reflow paragraphs: silently "correcting" a clause number
    or a chemical name in a hazard document would be worse than leaving it.
    """
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _RE_BAD_CHARS.sub("", text)
    text = text.replace("­", "")                 # soft hyphen
    text = text.replace(" ", " ")                # nbsp
    text = _RE_HYPHEN_BREAK.sub(r"\1\2", text)        # "equip-\nment" -> "equipment"

    kept = [ln for ln in text.split("\n") if not _RE_PAGE_NUM.match(ln)]
    text = "\n".join(kept)

    text = _RE_MANY_SPACES.sub(" ", text)
    text = _RE_TRAIL_SPACE.sub("", text)
    text = _RE_MANY_BLANKS.sub("\n\n", text)
    return text.strip()
